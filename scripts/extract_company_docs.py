"""Extract text from all company documents in Ai agentlar/."""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
SOURCE_DIR = ROOT / "Ai agentlar"
OUTPUT_DIR = ROOT / "scripts" / "extracted_docs"
SUPPORTED = {".docx", ".doc", ".txt", ".md", ".csv", ".json", ".xml", ".png", ".jpg", ".jpeg"}


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_image(path: Path) -> str:
    try:
        with Image.open(path) as img:
            return (
                f"[Image file: {path.name} | size={img.size} | mode={img.mode}] "
                "OCR not available — manual review required."
            )
    except Exception as exc:
        return f"[Image file unreadable: {path.name} | error={exc}]"


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_docx(path)
    if ext in {".txt", ".md", ".csv", ".json", ".xml"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if ext in {".png", ".jpg", ".jpeg"}:
        return extract_image(path)
    return f"[Unsupported extension {ext}]"


def main() -> int:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest: list[dict] = []

    if not SOURCE_DIR.is_dir():
        print(f"Source directory not found: {SOURCE_DIR}", file=sys.stderr)
        return 1

    for path in sorted(SOURCE_DIR.rglob("*")):
        if not path.is_file():
            continue
        rel = path.relative_to(SOURCE_DIR).as_posix()
        ext = path.suffix.lower()
        entry = {
            "relative_path": rel,
            "filename": path.name,
            "extension": ext,
            "size_bytes": path.stat().st_size,
        }
        try:
            if ext in SUPPORTED or ext == ".docx":
                text = extract_text(path)
                entry["char_count"] = len(text)
                entry["status"] = "ok" if text.strip() else "empty"
                out_file = OUTPUT_DIR / (rel.replace("/", "__") + ".txt")
                out_file.parent.mkdir(parents=True, exist_ok=True)
                out_file.write_text(text, encoding="utf-8")
                entry["extracted_to"] = str(out_file.relative_to(ROOT))
            else:
                entry["status"] = "unsupported"
                entry["char_count"] = 0
        except Exception as exc:
            entry["status"] = "error"
            entry["error"] = str(exc)
            entry["char_count"] = 0
        manifest.append(entry)

    manifest_path = OUTPUT_DIR / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps({"total_files": len(manifest), "manifest": str(manifest_path)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
